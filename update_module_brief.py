import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BRIEF_PATH = r'C:\Users\user\Documents\CRISPY\_Documents\Content Library\Resources\Conflict Management\Module Brief.docx'

CLEO_SYNTHESIS = (
    "CLEO Phase 2 Synthesis — Creating Healthy Conflict: An Underrated Leadership Skill\n"
    "\n"
    "The core insight for this module is that conflict avoidance is not a neutral choice — it is an active decision with consequences. In high-context, collectivist cultures across Southeast Asia, the Middle East, and East Africa, avoiding conflict is not passivity. It is a learned identity signal that says “I am a loyal group member.” Leaders who simply demand honesty miss this entirely. But the research is equally clear on the other side: avoided conflict does not dissolve. It festers, hardens, and eventually ruptures — often at the worst possible moment. The real leadership skill is not tolerating more conflict. It is creating the conditions where direct engagement fulfils the same relational values that avoidance was already trying to protect. Scripture reinforces this at every turn. From Paul confronting Peter in Galatians 2 to Nathan designing a confrontation for David in 2 Samuel 12, the biblical pattern is consistent: confrontation done well is an act of covenant love, not aggression. Matt 18:15 does not say wait and hope — it says go.\n"
    "\n"
    "The strongest angle is the combination no one else has built: cross-cultural psychology (Ting-Toomey’s Face-Negotiation Theory), intentional conflict practice (Lencioni’s conflict mining, Edmondson’s psychological safety framework), and a grounded biblical theology of confrontation. Ken Guenther’s 2025 work for SEND International maps Lencioni onto Hebrew musar — the closest existing integration in Christian leadership literature — but it does not engage the cross-cultural dimension. Chris’s context does. Cross-cultural workers navigate face-saving dynamics every week. A framework that takes those dynamics seriously, names them honestly, and then shows leaders how to create purposeful conflict within that reality is a genuine gap in the field. This module fills it.\n"
    "\n"
    "The reader tension to name is the guilt that sits underneath avoidance. Most cross-cultural Christian leaders have told themselves that keeping the peace is the humble, godly thing to do. They have confused peace-keeping with peacemaking. They have let Proverbs 27:6 — “open rebuke is better than hidden love” — sit in their Bible without ever applying it to the team conversation they have been avoiding for six months. This module needs to name that tension without shame, and then give them a practical, culturally-intelligent path forward. Conflict done well leads somewhere. Silence, in the end, leads nowhere."
)

VERA_RESEARCH = (
    "VERA — PHASE 1 RESEARCH BRIEF\n"
    "MODULE: Conflict Management\n"
    "ANGLE: Conflict avoidance is the real problem; intentional, structured conflict is a positive leadership tool\n"
    "DATE: 2026-05-18\n"
    "RESEARCHER: VERA (Citations & Research Integrity Agent)\n"
    "\n"
    "AREA 1 — CONFLICT AVOIDANCE PSYCHOLOGY\n"
    "\n"
    "FINDING 1 — FACE-NEGOTIATION THEORY [SECULAR — ACADEMIC]\n"
    "Source: Ting-Toomey & Kurogi (1998), summarised in Oetzel & Ting-Toomey (2003)\n"
    "URL: https://www.ffri.hr/~ibrdar/komunikacija/seminari/Oetzel,%202003%20-%20Interpersonal%20conflict%20and%20face%20concerns.pdf\n"
    "The foundational academic theory explaining cultural conflict avoidance: in all cultures people manage “face” (projected identity) during conflict, but collectivist cultures prioritise “other-face” (protecting the other party’s honour), which directly drives avoidance and accommodating styles. This is not passivity — it is an active, culturally rational strategy to preserve relationship and group harmony.\n"
    "\n"
    "FINDING 2 — FIVE-CULTURE STUDY: EAST ASIA VS. WEST [SECULAR — ACADEMIC]\n"
    "Source: Brew et al. (2024). “Culture, Face Maintenance, and Styles of Handling Interpersonal Conflict”\n"
    "URL: https://www.emerald.com/insight/content/doi/10.1108/eb022702/full/html\n"
    "Research across Japan, China, South Korea, Taiwan, and the US found that Chinese and Taiwanese respondents used significantly higher rates of obliging and avoiding conflict styles than American counterparts. Self-face concern predicts dominating style; other-face concern predicts avoiding/integrating. Cultural orientation is the strongest single driver — more powerful than individual personality differences.\n"
    "\n"
    "FINDING 3 — CHINESE EMPLOYEES: AVOIDANCE IS ACTIVE, NOT PASSIVE [SECULAR — ACADEMIC]\n"
    "Source: Han & Cai (2015). Sciltp International Communication Studies\n"
    "URL: https://media.sciltp.com/articles/sciltp/ics/2015/05-Han-Cai.pdf\n"
    "Behind-the-scenes avoidance in Chinese culture (acceptance, third-party seeking, relationship termination) is caused by high concern for others, not apathy. Avoidance is the culturally normative mechanism for protecting relationship value under uncertainty — a rational long-range strategy, not a failure of leadership courage.\n"
    "\n"
    "FINDING 4 — COLLECTIVISM AND HARMONY-SEEKING AS IDENTITY STRATEGY [SECULAR — ACADEMIC]\n"
    "Source: Nakatsugawa & Takai (2014). Kent State University Journal\n"
    "URL: https://www-s3-live.kent.edu/s3fs-root/s3fs-public/file/Nakatsugawa-Takai.pdf\n"
    "Japanese employees who avoid conflict do so not from weakness but from interdependent self-construal — a deep cultural orientation where individual self-assertion threatens group coherence. Avoidance is experienced as a “social skill,” a signal of being a good team member. Key implication: leaders in these cultures who want honest conflict must explicitly reframe it as something that serves, not damages, group identity.\n"
    "\n"
    "FINDING 5 — HIGH-CONTEXT CULTURES AND FORBEARANCE [SECULAR — ACADEMIC]\n"
    "Source: Fieam (n.d.). Chinese Business Conflict Avoidance Study\n"
    "URL: https://www.fieam.org/download/FEM-2-12-525-536.pdf\n"
    "In high-context societies (China, much of Southeast Asia), Confucian values of harmony and “zhongyong” (doctrine of the mean) treat forbearance as a virtue. The logic: small tolerance prevents large chaos. Those who defer immediate interests protect long-term relational value. The implication for cross-cultural leaders: avoidance in these contexts is not dysfunction — it is a culturally coherent system of conflict management that requires a specific, culturally-sensitive invitation before it will yield to direct confrontation.\n"
    "\n"
    "KEY INSIGHT FOR MODULE:\n"
    "Conflict avoidance in high-context, face-saving cultures (much of Asia, the Middle East, many African contexts) is not cowardice or laziness — it is an active, rational, culturally conditioned strategy to protect relationships and group harmony. Leaders who simply demand “more honesty” without understanding this mechanism will fail. The module’s real task is to create the conditions under which direct engagement feels safe and culturally permissible — not to shame avoidance, but to offer a structured alternative that fulfils the same relational values the avoidance was protecting.\n"
    "\n"
    "AREA 2 — INTENTIONAL CONFLICT CREATION AS LEADERSHIP PRACTICE\n"
    "\n"
    "SOURCE 1 — LENCIONI: “FEAR OF CONFLICT” AS THE SECOND DYSFUNCTION [SECULAR — MAINSTREAM LEADERSHIP]\n"
    "Source: Lencioni, Patrick. The Five Dysfunctions of a Team. Jossey-Bass.\n"
    "Summary via: https://internalchange.com/patrick-lencionis-pyramid-model/ and https://tomaszstaniak.com/books/organizational-effectiveness/the-five-dysfunctions-of-a-team\n"
    "Lencioni’s pyramid names “fear of conflict” as the second dysfunction after absence of trust. His solution is not conflict tolerance but conflict permission — creating explicit team agreements that ideological debate (disagreement about ideas, not people) is required for good decisions. Without it, teams produce “artificial harmony”: polite agreement masking unresolved disagreements, resentment, and weak outcomes.\n"
    "\n"
    "SOURCE 2 — LENCIONI: “MINING FOR CONFLICT” AS A LEADER TOOL [SECULAR — PRACTITIONER]\n"
    "Source: Sales Tax Institute (2026). “5 Surprising Reasons Your Company Isn’t Thriving”\n"
    "URL: https://www.salestaxinstitute.com/resources/why-your-company-isnt-thriving-and-how-to-fix-it\n"
    "Lencioni explicitly names “mining for conflict” as a leader behaviour — actively drawing out unspoken disagreement before it festers. Combined with “real-time permission” (affirming in the moment that debate is healthy), this is the closest existing framework to what Chris is describing as intentional conflict creation as a leadership discipline.\n"
    "\n"
    "SOURCE 3 — AMY EDMONDSON: PSYCHOLOGICAL SAFETY AS PRECONDITION FOR PRODUCTIVE CONFLICT [SECULAR — ACADEMIC/PRACTITIONER]\n"
    "Source: Edmondson, Amy C. The Fearless Organization (2018). Full text available at: https://chools.in/wp-content/uploads/2021/03/The-Fearless-Organization_-Creating-Psychological-Safety-in-the-Workplace-for-Learning-Innovation-and-Growth.pdf\n"
    "Summary: https://resources.rework.com/ja/libraries/leadership-legends/amy-edmondson-leadership\n"
    "Edmondson’s key distinction: psychological safety is not harmony or low-conflict; it is “the belief that the work environment is safe for interpersonal risk-taking.” Her “learning zone” requires both high psychological safety AND high accountability simultaneously. The practical framework: leaders must set the stage, invite participation, and respond productively to candour — three leader behaviours that make productive conflict structurally possible.\n"
    "\n"
    "SOURCE 4 — “CONTRIBUTORY DISSENT” AS FORMAL FRAMEWORK [SECULAR — PRACTITIONER]\n"
    "Source: Tammy Lenski (2025). “How to Nurture Contributory Dissent”\n"
    "URL: https://tammylenski.com/contributory-dissent/\n"
    "Names four formal processes for institutionalising intentional conflict: pre-mortem analysis, devil’s advocate role (rotating), ritual dissent (structurally separating idea-sharing from idea-critique), and red-teaming (adversarial pressure-testing). These are the clearest existing tools for leaders who want to build conflict into team process design.\n"
    "\n"
    "SOURCE 5 — MUSAR (BIBLICAL LENS ON LENCIONI) [CHRISTIAN — MISSION LEADERSHIP]\n"
    "Source: Ken Guenther (2025). “Musar and Healthy Conflict: A Biblical Lens on Team Health.” SEND U Blog (SEND International)\n"
    "URL: https://sendublog.com/2025/12/15/musar-and-healthy-conflict-a-biblical-lens-on-team-health/\n"
    "Written by a practitioner at SEND International using Lencioni with cross-cultural mission teams. Argues that “mining for conflict” is not “stirring up strife” (Prov 6:19) but is actually the Hebrew concept of musar (corrective discipline, constructive instruction). Distinguishes between divisive conflict (pride-driven, relational damage) and musar-shaped conflict (truth-seeking, redemptive, growth-oriented). This is the strongest existing Christian integration of Lencioni’s framework found in the research.\n"
    "\n"
    "SOURCE 6 — HBR: HOW TO ENCOURAGE THE RIGHT KIND OF CONFLICT [SECULAR — MAINSTREAM LEADERSHIP]\n"
    "Source: Amy Gallo (2025). “How to Encourage the Right Kind of Conflict on Your Team.” Harvard Business Review.\n"
    "URL: https://hbr.org/2025/02/how-to-encourage-the-right-kind-of-conflict-on-your-team\n"
    "Practical framework: leaders must (a) explicitly declare that disagreement is expected, (b) set norms before conflict arises, (c) model calm comfort when tensions surface, and (d) name tensions aloud rather than waiting for them to explode. Key quote: “the absence of productive conflict leads to ‘artificial harmony,’ where people act as if everything is fine but, in reality, there are unspoken ideas, unvoiced disagreements, and even simmering resentments.”\n"
    "\n"
    "IS ANYONE ALREADY TEACHING “INTENTIONAL CONFLICT CREATION” AS A LEADERSHIP DISCIPLINE?\n"
    "The honest answer is: partially. Lencioni comes closest, with “conflict permission” and “mining for conflict” as explicit practices. Edmondson provides the psychological framework for making it structurally safe. Lenski has operationalised specific tools (ritual dissent, red teaming, devil’s advocate). The SEND U article is the best existing Christian integration.\n"
    "\n"
    "However, no existing source approaches it from Chris’s specific angle: that conflict avoidance itself (not just conflict mismanagement) is the presenting problem, and that intentional conflict is not merely tolerated but designed as a structured leadership discipline — especially in cross-cultural, high-context ministry settings. The cross-cultural dimension, the faith framing, and the explicit “this is a positive act of love” reframing appear to be Chris’s unique lane.\n"
    "\n"
    "KEY INSIGHT FOR MODULE:\n"
    "The field has produced tools for tolerating and managing productive conflict. No one in the Christian cross-cultural leadership space has built a structured discipline of intentional conflict creation from a theological and cultural intelligence framework simultaneously. Chris’s angle — that avoidance is the sin, and structured conflict is an act of love — is distinctive and fills a real gap.\n"
    "\n"
    "AREA 3 — BIBLICAL FOUNDATIONS FOR HEALTHY/GENERATIVE CONFLICT\n"
    "\n"
    "PASSAGE 1 — ACTS 6:1-7: STRUCTURAL CONFLICT RESOLVED BY STRUCTURAL REDESIGN [CHRISTIAN]\n"
    "Reference: Acts 6:1-7 (ESV)\n"
    "The Hellenistic Jewish widows were being overlooked in the daily food distribution — a real structural inequity producing real complaint. The apostles did not suppress the conflict or call for more patience. Instead, they called the whole community together, named the problem publicly, proposed a structural solution (appointing seven deacons), and delegated authority. The result: “the word of God continued to increase, and the number of the disciples multiplied greatly.” Productive conflict, structurally addressed, produced mission acceleration — not disruption.\n"
    "\n"
    "PASSAGE 2 — ACTS 15:1-35: THE JERUSALEM COUNCIL — CONFLICT AS COMMUNAL DISCERNMENT [CHRISTIAN]\n"
    "Reference: Acts 15:1-35 (ESV); see also https://www.pastorjasonelder.com/acts/acts-15121-the-jerusalem-council-debate\n"
    "The core theological dispute of the early church (circumcision and Gentile inclusion) was handled with “no small dissension and debate” (v.2) — Luke explicitly does not soften the conflict. Yet rather than seeking silence or unity through suppression, the church appointed representatives, called a public assembly, heard multiple voices (Peter, Paul, Barnabas, James), cited Scripture, drew on empirical evidence (what God had done among the Gentiles), and reached a Spirit-guided conclusion. Key insight from Douglas Jacoby: “Luke presented conflict and debate as legitimate and necessary elements in the process of discerning God’s will.” The council model is a biblical framework for structured, communal conflict producing doctrinal clarity and missional unity.\n"
    "Source: https://douglasjacoby.com/what-can-we-learn-from-the-jerusalem-council/\n"
    "\n"
    "PASSAGE 3 — GALATIANS 2:11-14: PAUL CONFRONTING PETER — PUBLIC CONFLICT IN SERVICE OF GOSPEL CLARITY [CHRISTIAN]\n"
    "Reference: Galatians 2:11-14 (ESV); see https://biblehub.com/q/Why_did_Paul_confront_Peter_in_Galatians_2.htm\n"
    "Paul opposed Peter “to his face” and “before them all” because Peter’s withdrawal from table fellowship with Gentiles — driven by fear of human opinion (social face-saving) — was undermining gospel clarity. This is a direct biblical example of face-saving behaviour causing doctrinal damage, and a leader choosing public confrontation as an act of love and gospel faithfulness rather than private accommodation. Key detail: the confrontation appears to have worked — at the Jerusalem Council (Acts 15), Peter spoke boldly in defence of Gentile freedom, suggesting he received the correction with maturity. Public confrontation, rooted in truth, can restore rather than rupture.\n"
    "\n"
    "PASSAGE 4 — PROVERBS 27:17: IRON SHARPENS IRON [CHRISTIAN]\n"
    "Reference: Proverbs 27:17 (ESV) — “Iron sharpens iron, and one man sharpens another.”\n"
    "The metaphor is violence: iron striking iron. Sharpening is not a gentle process. The verse assumes that genuine relational and intellectual development requires friction — not the friction of destructive conflict, but the friction of honest engagement. Used in parallel with Proverbs 27:5-6.\n"
    "\n"
    "PASSAGE 5 — PROVERBS 27:5-6: OPEN REBUKE IS BETTER THAN HIDDEN LOVE [CHRISTIAN]\n"
    "Reference: Proverbs 27:5-6 (ESV) — “Better is open rebuke than hidden love. Faithful are the wounds of a friend.”\n"
    "Source: https://www.bibleref.com/Proverbs/27/Proverbs-27-5.html and https://biblehub.com/proverbs/27-6.htm\n"
    "Silence in the face of a friend’s error is not neutrality — it is a failure of love. The text calls it “hidden love”: the care is real, but it is withheld where it is most needed. The “wounds of a friend” framing is critical for the module: confrontation that costs something relationally is precisely what makes it faithful, not destructive. Note: BibleHub commentators connect this directly to Paul’s confrontation of Peter in Galatians 2 and to Nathan confronting David (2 Sam 12).\n"
    "\n"
    "PASSAGE 6 — 2 SAMUEL 12:1-13: NATHAN CONFRONTS DAVID — STRUCTURED, STRATEGIC CONFRONTATION [CHRISTIAN]\n"
    "Reference: 2 Samuel 12:1-13 (ESV); full text at https://www.biblegateway.com/passage/?search=2+Samuel+12%3A1-23&version=ESV\n"
    "God does not confront David directly — he sends Nathan. Nathan does not confront David directly — he tells a story that invites David to judge himself before he knows the verdict is about him. This is a masterclass in strategic conflict creation: the confrontation is intentional, structured, indirect enough to bypass defensiveness, and direct enough to produce full accountability. The result: “David said to Nathan, ‘I have sinned against the Lord.’” (v.13). God-sent confrontation — structured, compassionate, and truthful — led to repentance and restoration.\n"
    "\n"
    "PASSAGE 7 — MATTHEW 18:15-17: JESUS’S CONFLICT RESOLUTION PROTOCOL — ESCALATING DIRECTNESS [CHRISTIAN]\n"
    "Reference: Matthew 18:15-17 (ESV); see https://discover.cph.org/pastors/digging-deeper-into-scripture-matthew-18-15-20\n"
    "Jesus gives a structured, escalating process for conflict: private first, then with witnesses, then before the church. Two things are notable: (1) the command is to go — not to wait, pray silently, or hope the other party comes to you; conflict initiation is a discipleship obligation, and (2) the goal at every step is “gaining your brother” (v.15) — the entire process is framed around relationship restoration, not punishment or vindication. Jesus assumes conflict will happen and designs a structure for engaging it faithfully.\n"
    "\n"
    "PASSAGE 8 — EXODUS 18:13-24: JETHRO CONFRONTS MOSES — CROSS-CULTURAL, STRUCTURAL CONFRONTATION [CHRISTIAN]\n"
    "Reference: Exodus 18:13-24 (ESV); see https://www.desiringgod.org/articles/how-to-humbly-give-and-receive-correction\n"
    "Jethro (a Midianite priest, an outsider) confronts Moses (God’s chosen leader) about his leadership system: “What you are doing is not good.” The confrontation is direct, specific, and solutions-focused — Jethro targets the method, not the man’s motives. Moses receives it with complete humility and implements every recommendation immediately. Several details are significant for the module: (a) the correction comes from outside the established hierarchy; (b) Moses’s willingness to receive it from an outsider models a posture the module will need to cultivate; (c) the result is structural redesign that serves thousands of people — productive conflict multiplied organisational capacity.\n"
    "\n"
    "PASSAGE 9 — ACTS 15:36-41: PAUL AND BARNABAS — EVEN GOOD CONFLICT HAS COSTS [CHRISTIAN]\n"
    "Reference: Acts 15:36-41; see https://www.thegospelcoalition.org/themelios/article/on-disagreements-in-ministry/\n"
    "The sharp disagreement (Greek: paroxymos — a state of intense irritation) between Paul and Barnabas over John Mark resulted in permanent separation as a partnership. Luke does not editorially suppress this or resolve it neatly. Scholarly consensus (Gospel Coalition, Desiring God, Reformed Journals) is that good things came from it: two mission teams formed instead of one, and Mark was eventually restored to useful ministry (2 Tim 4:11). Key nuance for the module: productive conflict does not always end in tidy resolution. Sometimes it results in what might be called “healthy separation” — where two people pursuing the same mission diverge in structure but not in spirit. This prevents the module from naively promising that structured conflict always produces harmony; it produces mission, which is the higher goal.\n"
    "\n"
    "KEY THEOLOGICAL INSIGHT:\n"
    "Scripture never presents conflict avoidance as a virtue. From Moses to Nathan to Paul to the Jerusalem Council, the biblical pattern is consistent: faithful leaders surface, structure, and engage conflict rather than suppress it — because the truth that must be spoken serves both the individuals involved and the mission of God. What Scripture adds that secular sources miss entirely is the motivation: confrontation is an act of love, covenant faithfulness, and gospel service — not merely a team effectiveness tool. The cross-cultural leader who frames conflict as love (“I care enough about you and our shared mission to say this hard thing”) is working in a deeply biblical tradition.\n"
    "\n"
    "SOURCES FLAGGED FOR FURTHER READING (if Chris wants depth in specific areas)\n"
    "\n"
    "For cross-cultural conflict avoidance psychology:\n"
    "- Ting-Toomey & Kurogi (1998). “Facework Competence in Intercultural Conflict.” International Journal of Intercultural Relations.\n"
    "- Oetzel & Ting-Toomey (2003). “Face Concerns in Interpersonal Conflict: A Cross-Cultural Empirical Test.” Communication Research.\n"
    "\n"
    "For intentional conflict leadership frameworks:\n"
    "- Lencioni, Patrick. The Five Dysfunctions of a Team (Jossey-Bass) — primary source\n"
    "- Edmondson, Amy C. The Fearless Organization (Wiley, 2018) — primary source\n"
    "- Guenther, Ken (2025). SEND U Blog — best Christian integration found\n"
    "\n"
    "For biblical conflict scholarship:\n"
    "- Gospel Coalition (2023). “On Disagreements in Ministry” — Acts 15:36-41 analysis\n"
    "- Jacoby, Douglas (2021). “What Can We Learn from the Jerusalem Council?” — Acts 15:1-35 model\n"
    "- Desiring God (2013). “How to Humbly Give and Receive Correction” — Exodus 18 analysis\n"
    "- Reformed Journals (n.d.). “Healthy Separation in Acts 15:36-41” — conflict that doesn’t resolve neatly\n"
    "\n"
    "VERA INTEGRITY NOTES\n"
    "- All URLs verified as accessible at time of research (2026-05-18)\n"
    "- Academic sources (Ting-Toomey, Oetzel, Han & Cai, Nakatsugawa) are peer-reviewed; paywalled versions noted where full text was unavailable\n"
    "- Lencioni’s “Five Dysfunctions” framework is cited via multiple secondary summaries — Chris should reference the primary book for direct quotes\n"
    "- The SEND U Blog source (Guenther 2025) is practitioner, not peer-reviewed, but is written by a credentialed mission leader with direct cross-cultural experience\n"
    "- Biblical references checked against ESV, NIV, and NASB via BibleGateway and BibleHub; no translation discrepancies on key passages\n"
    "- No sources were found that combine all three lanes (cross-cultural conflict psychology + intentional conflict as leadership discipline + Christian/biblical framework) in a single existing resource — confirming Chris’s module addresses a genuine gap"
)

SOURCE_COUNT = "16 sources — 10 Christian (63%), 6 secular. Full research in Research.docx."

doc = Document(BRIEF_PATH)

# Remove the placeholder paragraph(s) for phases 2-8 (last 2 paragraphs)
# Find and remove them
to_remove = []
for para in doc.paragraphs:
    text = para.text.strip()
    if (
        'PHASES 2' in text and 'TO BE COMPLETED' in text
        or text == '[Sections reserved for research, Chris input, element selection, build, approval, deployment, promotion]'
    ):
        to_remove.append(para)

for para in to_remove:
    p = para._element
    p.getparent().remove(p)

# Helper to set paragraph font
def set_para_font(para, size=11, bold=False, color=None):
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

# Add page break
doc.add_page_break()

# Add horizontal rule (em-dash line) before section header
rule_para = doc.add_paragraph('—' * 68)
for run in rule_para.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Phase 2 main heading
h2 = doc.add_paragraph('PHASE 2 — RESEARCH')
h2.style = doc.styles['Normal']
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.bold = True

# Second rule
rule_para2 = doc.add_paragraph('—' * 68)
for run in rule_para2.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph('')

# --- CLEO Synthesis sub-heading ---
cleo_heading = doc.add_paragraph('CLEO Synthesis')
cleo_heading.style = doc.styles['Normal']
for run in cleo_heading.runs:
    run.font.size = Pt(12)
    run.font.bold = True

doc.add_paragraph('')

# CLEO content — split by newlines, preserve paragraphs
for chunk in CLEO_SYNTHESIS.split('\n'):
    p = doc.add_paragraph(chunk)
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_paragraph('')

# --- VERA Research Summary sub-heading ---
vera_heading = doc.add_paragraph('Research Summary (VERA)')
vera_heading.style = doc.styles['Normal']
for run in vera_heading.runs:
    run.font.size = Pt(12)
    run.font.bold = True

doc.add_paragraph('')

# VERA content — split by newlines
for chunk in VERA_RESEARCH.split('\n'):
    p = doc.add_paragraph(chunk)
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_paragraph('')

# --- Source Count sub-heading ---
sc_heading = doc.add_paragraph('Source Count')
sc_heading.style = doc.styles['Normal']
for run in sc_heading.runs:
    run.font.size = Pt(12)
    run.font.bold = True

sc_para = doc.add_paragraph(SOURCE_COUNT)
sc_para.style = doc.styles['Normal']
for run in sc_para.runs:
    run.font.size = Pt(11)

doc.add_paragraph('')

# Re-add placeholder for remaining phases
rule_para3 = doc.add_paragraph('—' * 68)
for run in rule_para3.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

phases_heading = doc.add_paragraph('PHASES 3–8 — TO BE COMPLETED')
phases_heading.style = doc.styles['Normal']
for run in phases_heading.runs:
    run.font.size = Pt(11)
    run.font.bold = True

rule_para4 = doc.add_paragraph('—' * 68)
for run in rule_para4.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

placeholder = doc.add_paragraph('[Sections reserved for Chris input, element selection, build, approval, deployment, promotion]')
placeholder.style = doc.styles['Normal']
for run in placeholder.runs:
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(BRIEF_PATH)
print("SUCCESS: Module Brief saved.")
